import io
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from utils.report_image_embedder import (
    _append_tag_description,
    embed_images_into_docx,
    extract_angle_bracket_placeholders,
)


@pytest.mark.unit
class TestReportImageEmbedderUnit:
    @patch("utils.report_image_embedder._generate_paragraph_from_tags")
    def test_append_tag_description_summary_and_raw(self, mock_summary):
        mock_summary.return_value = "Damage was observed on the container door."
        doc = Document()
        _append_tag_description(doc, "rust, dent (NEEDS ATTENTION)")

        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "Damage was observed on the container door."
        assert not doc.paragraphs[0].runs[0].italic
        assert doc.paragraphs[1].text == "rust, dent (NEEDS ATTENTION)"
        assert doc.paragraphs[1].runs[0].italic

    @patch("utils.report_image_embedder._generate_paragraph_from_tags")
    def test_append_tag_description_raw_only_when_llm_fails(self, mock_summary):
        mock_summary.return_value = None
        doc = Document()
        _append_tag_description(doc, "clean, no damage")

        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "clean, no damage"
        assert doc.paragraphs[0].runs[0].italic

    def test_generate_paragraph_from_tags_returns_llm_text(self):
        from utils.report_image_embedder import _generate_paragraph_from_tags

        mock_response = MagicMock()
        mock_response.text = "  Rust was noted on the door panel.  "
        with patch("utils.report_image_embedder.getLLMClient") as mock_client, patch(
            "utils.report_image_embedder.gemini_utils.prompt_text",
            return_value=(mock_response, None),
        ):
            mock_client.return_value = MagicMock()
            result = _generate_paragraph_from_tags("rust, dent")

        assert result == "Rust was noted on the door panel."

    @patch("utils.report_image_embedder._prefetch_tag_summaries")
    def test_embed_appends_tags_when_picture_fails(self, mock_prefetch):
        mock_prefetch.return_value = {"rust": "Rust was noted."}
        doc = Document()
        doc.add_paragraph("<photo.jpg>")
        buf = io.BytesIO()
        doc.save(buf)

        out_bytes, placeholders = embed_images_into_docx(
            buf.getvalue(),
            [("photo.jpg", b"not-a-real-image", "rust")],
        )

        out_doc = Document(io.BytesIO(out_bytes))
        joined = "\n".join(p.text for p in out_doc.paragraphs)
        assert "Unable to embed image" in joined
        assert "rust" in joined
        assert placeholders == []

    @patch("utils.report_image_embedder._prefetch_tag_summaries", return_value={})
    def test_embed_returns_remaining_placeholders_only(self, _mock_prefetch):
        doc = Document()
        doc.add_paragraph("<keep.jpg>")
        doc.add_paragraph("<strip.jpg>")
        buf = io.BytesIO()
        doc.save(buf)

        out_bytes, placeholders = embed_images_into_docx(
            buf.getvalue(),
            [("strip.jpg", b"\x89PNG\r\n\x1a\n", None)],
        )

        assert placeholders == ["keep.jpg"]
        out_doc = Document(io.BytesIO(out_bytes))
        remaining = extract_angle_bracket_placeholders(out_doc)
        assert remaining == ["keep.jpg"]
