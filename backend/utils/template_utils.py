# utils/template_utils.py
import re
import json
from docx import Document
from docx.shared import Pt
from io import BytesIO
from typing import Optional


def extract_needs_attention_images(tags_json: str) -> list[str]:
    """Extract image names that have NEEDS ATTENTION tags."""
    try:
        tags = json.loads(tags_json)
        flagged = []
        for item in tags:
            image_name = item.get("image_name", "")
            item_tags = item.get("tags", [])
            if any("NEEDS ATTENTION" in tag for tag in item_tags):
                flagged.append(image_name)
        return flagged
    except Exception:
        return []


def fill_template(template_path: str, filled_values: dict, tags_json: Optional[str] = None) -> BytesIO:
    doc = Document(template_path)

    def replace_in_paragraph(para, values):
        # check full paragraph text first
        full_text = para.text
        if not any(f"<survey_ops_plus:{key}>" in full_text for key in values):
            return
        
        # rebuild the paragraph text with replacements
        new_text = full_text
        for key, value in values.items():
            placeholder = f"<survey_ops_plus:{key}>"
            new_text = new_text.replace(placeholder, value or "N/A")
        
        # clear all runs and set text in first run
        if new_text != full_text:
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_text
                else:
                    run.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para, filled_values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, filled_values)

    if tags_json:
        flagged_images = extract_needs_attention_images(tags_json)
        if flagged_images:
            doc.add_paragraph("")
            p = doc.add_paragraph("Images Requiring Attention:")
            p.runs[0].bold = True
            for img_name in flagged_images:
                doc.add_paragraph(f"<{img_name}>")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_template_tags(template_path: str) -> list[str]:
    """Extract all <tag> placeholders from a .docx template."""
    doc = Document(template_path)
    tags = set()
    pattern = re.compile(r'<survey_ops_plus:([^>]+)>')

    for para in doc.paragraphs:
        for match in pattern.finditer(para.text):
            tags.add(match.group(1))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in pattern.finditer(para.text):
                        tags.add(match.group(1))
    
    return list(tags)
