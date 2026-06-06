"""Append survey images to the end of a .docx and strip matching <name> placeholder paragraphs."""
from __future__ import annotations

import io
import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.shared import Inches, Pt

from llm.gemini import utils as gemini_utils
from llm.prompts.prompts import image_tags_summary_prompt
from utils.llm_client import getLLMClient, resolve_google_model

logger = logging.getLogger(__name__)

SECTION_TITLE = "Survey image attachments"
_PLACEHOLDER_LINE = re.compile(r"^<([^>]+)>$")


def _delete_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_angle_bracket_placeholders(doc: Document) -> list[str]:
    found: set[str] = set()
    for p in _iter_all_paragraphs(doc):
        t = p.text.strip()
        m = _PLACEHOLDER_LINE.match(t)
        if m:
            found.add(m.group(1).strip())
    return sorted(found)


def strip_embedded_image_appendix(doc: Document) -> None:
    """Remove a prior appendix starting at the fixed section title through end of body."""
    while True:
        start_idx: int | None = None
        paras = list(doc.paragraphs)
        for i, p in enumerate(paras):
            if p.text.strip() == SECTION_TITLE:
                start_idx = i
                break
        if start_idx is None:
            return
        for j in range(len(doc.paragraphs) - 1, start_idx - 1, -1):
            _delete_paragraph(doc.paragraphs[j])


def remove_placeholder_paragraphs_for_names(doc: Document, names: set[str]) -> None:
    if not names:
        return
    for _ in range(2000):
        removed = False
        for p in list(doc.paragraphs):
            t = p.text.strip()
            m = _PLACEHOLDER_LINE.match(t)
            if m and m.group(1).strip() in names:
                _delete_paragraph(p)
                removed = True
                break
        if not removed:
            break


def _generate_paragraph_from_tags(tags_str: str) -> str | None:
    tags_str = tags_str.strip()
    if not tags_str:
        return None
    try:
        client = getLLMClient("google")
        model_name = resolve_google_model()
        prompt = image_tags_summary_prompt(tags_str)
        response, err = gemini_utils.prompt_text(client, prompt, model_name)
        if err or response is None or not response.text:
            logger.warning("Tag summary LLM failed: %s", err)
            return None
        return response.text.strip()
    except Exception as e:
        logger.warning("Tag summary LLM exception: %s", e)
        return None


def _prefetch_tag_summaries(
    images_ordered: list[tuple[str, bytes, str | None]],
) -> dict[str, str | None]:
    descriptions = {
        description.strip()
        for _, _, description in images_ordered
        if description and description.strip()
    }
    if not descriptions:
        return {}

    summaries: dict[str, str | None] = {}
    workers = min(4, len(descriptions))
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = {
            executor.submit(_generate_paragraph_from_tags, desc): desc
            for desc in descriptions
        }
        for future in as_completed(futures):
            desc = futures[future]
            try:
                summaries[desc] = future.result()
            except Exception as e:
                logger.warning("Tag summary prefetch failed for %s: %s", desc, e)
                summaries[desc] = None
    return summaries


def _append_tag_description(
    doc: Document,
    tags_str: str,
    summary: str | None = None,
) -> None:
    tags_str = tags_str.strip()
    if not tags_str:
        return

    if summary is None:
        summary = _generate_paragraph_from_tags(tags_str)
    if summary:
        p_summary = doc.add_paragraph()
        r_summary = p_summary.add_run(summary)
        r_summary.font.size = Pt(10)

    p_raw = doc.add_paragraph()
    r_raw = p_raw.add_run(tags_str)
    r_raw.font.size = Pt(9)
    r_raw.italic = True


def embed_images_into_docx(
    doc_bytes: bytes,
    images_ordered: list[tuple[str, bytes, str | None]],
) -> tuple[bytes, list[str]]:
    """
    Returns (new_doc_bytes, angle-bracket placeholder names still in the document).
    """
    doc = Document(io.BytesIO(doc_bytes))
    names_to_strip = {item[0].strip() for item in images_ordered if item[0].strip()}
    remove_placeholder_paragraphs_for_names(doc, names_to_strip)
    strip_embedded_image_appendix(doc)

    if images_ordered:
        tag_summaries = _prefetch_tag_summaries(images_ordered)
        p_title = doc.add_paragraph(SECTION_TITLE)
        p_title.runs[0].bold = True
        for i, (display_name, img_bytes, description) in enumerate(images_ordered):
            if i > 0:
                doc.add_paragraph("")
            head = doc.add_paragraph()
            r = head.add_run(display_name)
            r.bold = True
            r.font.size = Pt(11)
            try:
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
            except Exception:
                logger.warning(
                    "Failed to embed image bytes for %s", display_name, exc_info=True
                )
                doc.add_paragraph(f"(Unable to embed image: {display_name})")

            desc_text = (description or "").strip()
            if desc_text:
                try:
                    _append_tag_description(
                        doc,
                        desc_text,
                        tag_summaries.get(desc_text),
                    )
                except Exception:
                    logger.warning(
                        "Failed to append tag description for %s",
                        display_name,
                        exc_info=True,
                    )

    placeholders_remaining = extract_angle_bracket_placeholders(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), placeholders_remaining
